from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches
headers = {
    'authority': 'app.****.com',
    'cache-control': 'max-age=0',
    'sec-ch-ua': '"Google Chrome";v="89", "Chromium";v="89", ";Not A Brand";v="99"',
    'sec-ch-ua-mobile': '?0',
    'upgrade-insecure-requests': '1',
    'origin': '***url***',
    'content-type': 'application/x-www-form-urlencoded',
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.90 Safari/537.36',
    'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
    'sec-fetch-site': 'same-origin',
    'sec-fetch-mode': 'navigate',
    'sec-fetch-user': '?1',
    'sec-fetch-dest': 'document',
    'referer': '***url login***',
    'accept-language': 'es-ES,es;q=0.9,en;q=0.8,ca;q=0.7',
    'cookie': '*****',
}

data = {
  'email': '***@***.es',
  'password': '*****',
  'rememberme': 'on'
}
with requests.Session() as s:
    response = s.post('***url***', headers=headers, data=data)
    data = s.get('***url página***')
    texto = data.text.split('<small>#')
    lista = list()
    for i in texto:
        lista.append(i[0:6])
    lista = lista[1:]
    print('Elementos: ', len(lista))

    campaign = list();

    for i in lista[0:len(lista)]:
        request = s.get('***url+id****'+i+'#Messages')
        ndata = request.text
        soup = BeautifulSoup(ndata, 'html.parser')
        title = soup.find('span', class_='campaign-title')
        title = title.text
        title = i + " Activa "+ title
        questions = [element.text for element in soup.findAll('textarea')]

        document = Document()
        document.add_paragraph(title)
        for i in questions:
            document.add_paragraph(i, style='ListBullet')
        title = title.replace('/','')
        title = title.replace('?','-')
        document.save(title+'.docx')







